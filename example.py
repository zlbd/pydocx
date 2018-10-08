#!/usr/bin/env python
#-*-coding:utf-8-*-
# Example from: https://python-docx.readthedocs.io/en/latest/
# Author: zlbd

from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 2)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

document.add_picture('img/monty-truth.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

paragraph = document.add_paragraph(u'明天')
run = paragraph.add_run(u'会')
run.bold = True
paragraph.add_run(u' 更好.')
paragraph.add_run('dolor').bold = True
run = paragraph.add_run('dolor')
run.bold = True
run.underline = True
paragraph = document.add_paragraph('Normal text, ')
paragraph.add_run('text with emphasis.', 'Emphasis')

document.save('demo.docx')
print("Finished 1!")

document = Document()
document.save(u'测试.docx')
print("Finished 2!")

document = Document('demo.docx')
paragraph = document.add_paragraph()
paragraph.add_run('This is demo 2')
document.save('demo2.docx')
print("Finished 3!")
