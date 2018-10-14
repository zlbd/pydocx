#!/usr/bin/env python
#-*- coding:utf-8 -*-
# Example from: https://www.jianshu.com/p/1f60cdd9655
import docx
from docx          import Document
from docx.shared   import *
from docx.oxml.ns  import qn
from docx.enum.dml import *
from docx.enum.style    import *
from docx.enum.section  import WD_SECTION 
#from docx.enum.section  import WD_UNDERLINE
import sys
reload(sys)
sys.setdefaultencoding('utf-8')

def main():
    # Accessing sections
    print("\n##### Accessing sections")
    document = Document()
    sections = document.sections
    print sections
    print len(sections)
    section = sections[0]
    print section
    for section in sections:
        print(section.start_type)
        print(section.header_distance.pt)
    # Add a new section
    print("\n#### Add a new section")
    current_section = document.sections[-1]
    print(current_section.start_type)
    new_section = document.add_section(WD_SECTION.ODD_PAGE)
    print(new_section.start_type)
    new_section.start_type = WD_SECTION.NEW_PAGE
    print(new_section.start_type)
    #print(current_section.header)
    # Section dimensions and orientation
    print(section.orientation, section.page_width.inches, section.page_height.pt)
    # Page margins
    print("\n#### Page margins")
    print(section.left_margin.pt, section.right_margin.pt)
    print(section.top_margin.pt, section.bottom_margin.pt)
    print(section.gutter)
    print(section.header_distance.pt, section.footer_distance.pt)
    section.left_margin, section.right_margin = Pt(80.0), Pt(80.0)
    print(section.left_margin.pt, section.right_margin.pt)
    #Access a style 
    print("\n#### Access a style")
    styles = document.styles
    paragraph_styles = [
        s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH
    ]
    for style in paragraph_styles:
        print(style.name)
    print(len(styles), len(paragraph_styles))
    # Apply a style
    print("\n#### Apply a style")
    paragraph = document.add_paragraph()
    print(paragraph.style.name)
    paragraph.style = document.styles['Heading 1']
    print(paragraph.style.name)
    paragraph.style = 'List Bullet'
    print(paragraph.style.name)
    paragraph = document.add_paragraph(style = 'Body Text')
    print(paragraph.style.name)
    body_test_style = document.styles['Body Text']
    paragraph = document.add_paragraph(style = body_test_style)
    print(type(body_test_style))
    print(paragraph.style.name)
    # Add or delete a style
    print("\n#### Add or delete a style")
    styles = document.styles
    style = styles.add_style('Citation', WD_STYLE_TYPE.PARAGRAPH)
    print(style.name, style.type)
    print(len(styles))
    styles['Citation'].delete()
    print(len(styles))
    # #### Define character formatting
    print("\n#### Define character formatting")
    style = document.styles['Normal']
    font  = style.font
    print(font.name, font.size)
    font.name = 'Calibri'
    font.size = Pt(12)
    print(font.name, font.size.pt)
    print(font.bold, font.italic)
    #font.underline = WD_UNDERLINE.DOT_DASH
    ##### Define paragraph formatting
    print("\n#### Define paragraph formatting")
    sytle = document.styles.add_style('Indent', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style.paragraph_format
    paragraph_format.left_indent = Inches(0.25)
    paragraph_format.first_line_indent = Inches(-0.25)
    paragraph_format.space_before = Pt(12)
    paragraph_format.widow_control= True
    #Use paragraph-specific sytle properties
    print("\n#### Use paragraph-specific sytle properties")
    sytles = document.styles
    styles['Heading 1'].next_paragraph_style = styles['Body Text']
    heading_1_style = styles['Heading 1']
    print(heading_1_style.next_paragraph_style.name)
    heading_1_style.next_paragraph_style = heading_1_style
    #Working with Latent Styles
    print("\n#####Working with Latent Styles")
    latent_styles = document.styles.latent_styles
    print(len(latent_styles))
    latent_style_names = [ls.name for ls in latent_styles]
    #print(latent_style_names)
    latent_quote = latent_styles['Quote']
    print(latent_quote, latent_quote.priority)
    #my_laytent_style = docx.styles.latent._LatentStyle
    #print(my_laytent_style)






if __name__ == "__main__":
    main()

