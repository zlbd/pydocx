#!/usr/bin/env python
#-*- coding:utf-8 -*-
# Example from: https://www.jianshu.com/p/1f60cdd9655
from docx          import Document
from docx.shared   import *
from docx.oxml.ns  import qn
from docx.enum.dml import MSO_THEME_COLOR

import sys
reload(sys)
sys.setdefaultencoding('utf-8')

def add_test_paragraph(paragraph):
    #设置下划线
    run = paragraph.add_run(u'设置下划线')
    run.font.underline = True

    #设置颜色
    run = paragraph.add_run(u'设置颜色')
    run.font.color.rgb = RGBColor(0xFF, 0x24, 0xE9)
    print run.text, run.font.color.rgb[0], run.font.color.rgb[1]

    #设置主题颜色
    arrColor = [
        MSO_THEME_COLOR.ACCENT_1,
        MSO_THEME_COLOR.ACCENT_2,
        MSO_THEME_COLOR.ACCENT_3,
        MSO_THEME_COLOR.ACCENT_4,
        MSO_THEME_COLOR.ACCENT_5,
        MSO_THEME_COLOR.ACCENT_6
    ]
    for color in arrColor:
        run = paragraph.add_run(u'\n设置颜色theme_color' + str(color))
        run.font.color.theme_color = color
    print run.text, run.font.color.rgb[0], run.font.color.rgb[1]

def main():
    #打开文档
    document = Document()

    #加入不同等级的标题
    document.add_heading('Document Title',0)
    document.add_heading(u'二级标题',1)
    document.add_heading(u'二级标题',2)

    #添加文本
    paragraph = document.add_paragraph(u'添加了文本')
    #设置字号
    run = paragraph.add_run(u'设置字号')
    run.font.size=Pt(24)

    #设置字体
    run = paragraph.add_run('Set Font,')
    run.font.name='Consolas'

    #设置中文字体
    run = paragraph.add_run(u'设置中文字体，')
    run.font.name=u'宋体'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    #设置斜体
    run = paragraph.add_run(u'斜体、')
    run.italic = True

    #设置粗体
    run = paragraph.add_run(u'粗体').bold = True

    #自定义测试
    add_test_paragraph(document.add_paragraph())

    #增加引用
    document.add_paragraph('Intense quote', style='Intense Quote')

    #增加有序列表
    document.add_paragraph(
        u'有序列表元素1',style='List Number'
    )
    document.add_paragraph(
        u'有序列别元素2',style='List Number'
    )

    #增加无序列表
    document.add_paragraph(
        u'无序列表元素1',style='List Bullet'
    )
    document.add_paragraph(
        u'无序列表元素2',style='List Bullet'
    )

    #增加图片（此处使用相对位置）
    document.add_picture('img/monty-truth.png', width=Inches(1.25))

    #增加表格
    table = document.add_table(rows=3,cols=3)
    hdr_cells=table.rows[0].cells
    hdr_cells[0].text=u"第一列"
    hdr_cells[1].text=u"第二列"
    hdr_cells[2].text=u"第三列"

    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = '2'
    hdr_cells[1].text = 'aerszvfdgx'
    hdr_cells[2].text = 'abdzfgxfdf'

    hdr_cells = table.rows[2].cells
    hdr_cells[0].text = '3'
    hdr_cells[1].text = 'cafdwvaef'
    hdr_cells[2].text = 'aabs zfgf'

    #增加分页
    document.add_page_break()

    #保存文件
    document.save('demo.docx')
    print('demo.docx is saved OK!')

if __name__ == "__main__":
    main()

